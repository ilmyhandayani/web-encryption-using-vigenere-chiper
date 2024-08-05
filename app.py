from flask import Flask, request, render_template, send_file, redirect, url_for, session
import os
from docx import Document
import mysql.connector

app = Flask(__name__)
app.secret_key = 'ilmy24'

# Konfigurasi koneksi database MySQL
db = mysql.connector.connect(
    host="localhost",
    user="root",
    password="",
    database="kriptografi"
)

cursor = db.cursor()

@app.route('/', methods=['GET', 'POST'])
def home():
    if 'username' in session:
        if request.method == 'POST':
            # Handle file encryption
            input_file = request.files['file']
            key = request.form['key']

            # Save the input file
            input_file_path = os.path.join(app.root_path, 'input.txt') if input_file.filename.endswith('.txt') else os.path.join(app.root_path, 'input.docx')
            input_file.save(input_file_path)

            # Encrypt the file
            encrypted_file = os.path.join(app.root_path, 'encrypted.txt') if input_file.filename.endswith('.txt') else os.path.join(app.root_path, 'encrypted.docx')
            encrypt_file(input_file_path, encrypted_file, key)

            # Read the input and encrypted files as binary data
            with open(input_file_path, 'rb') as f:
                input_file_data = f.read()
            with open(encrypted_file, 'rb') as f:
                encrypted_file_data = f.read()

            # Store the file data in the database
            username = session['username']
            update_query = "UPDATE users SET input_file = %s, encrypted_file = %s WHERE username = %s"
            cursor.execute(update_query, (input_file_data, encrypted_file_data, username))
            db.commit()

            # Delete the input and encrypted files
            # # os.remove(input_file_path)
            # # os.remove(encrypted_file)

            return render_template('index.html', message='File encrypted successfully.', encrypted_file=encrypted_file)

        return render_template('index.html')

    return redirect(url_for('login'))

@app.route('/decrypt', methods=['POST'])
def decrypt():
    # Handle file decryption
    encrypted_file = request.files['file']
    key = request.form['key']

    # Save the encrypted file
    encrypted_file_path = 'encrypted.txt' if encrypted_file.filename.endswith('.txt') else 'encrypted.docx'
    encrypted_file.save(encrypted_file_path)

    # Decrypt the file
    decrypted_file = 'decrypted.txt' if encrypted_file.filename.endswith('.txt') else 'decrypted.docx'
    decrypt_file(encrypted_file_path, decrypted_file, key)

    # Delete the encrypted file
    os.remove(encrypted_file_path)

    return render_template('index.html', message='File decrypted successfully.', decrypted_file=decrypted_file)

@app.route('/download/<path:filename>', methods=['GET'])
def download(filename):
    return send_file(filename, as_attachment=True)

def encrypt_file(input_file, output_file, key):
    if input_file.endswith('.txt'):
        with open(input_file, 'r') as file:
            text = file.read()

        encrypted_text = encrypt_text(text, key)

        with open(output_file, 'w') as file:
            file.write(encrypted_text)
    elif input_file.endswith('.docx'):
        doc = Document(input_file)

        for paragraph in doc.paragraphs:
            encrypted_text = encrypt_text(paragraph.text, key)
            paragraph.text = encrypted_text

        doc.save(output_file)

def decrypt_file(input_file, output_file, key):
    if input_file.endswith('.txt'):
        with open(input_file, 'r') as file:
            encrypted_text = file.read()

        decrypted_text = decrypt_text(encrypted_text, key)

        with open(output_file, 'w') as file:
            file.write(decrypted_text)
    elif input_file.endswith('.docx'):
        doc = Document(input_file)

        for paragraph in doc.paragraphs:
            decrypted_text = decrypt_text(paragraph.text, key)
            paragraph.text = decrypted_text

        doc.save(output_file)

def encrypt_text(text, key):
    encrypted_text = text[::-1]  # Reverse the text
    return encrypted_text

def decrypt_text(text, key):
    decrypted_text = text[::-1]  # Reverse the text
    return decrypted_text

@app.route('/login', methods=['GET', 'POST'])
def login():
    if 'username' in session:
        return redirect(url_for('home'))

    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        # Query untuk memeriksa keberadaan username dan password di database
        query = "SELECT * FROM users WHERE username = %s AND password = %s"
        cursor.execute(query, (username, password))
        user = cursor.fetchone()

        if user:
            session['username'] = user[1]  # Menggunakan session dengan username dari database
            return redirect(url_for('home'))
        else:
            return render_template('login.html', error='Invalid username or password')

    return render_template('login.html')

@app.route('/logout', methods=['GET', 'POST'])
def logout():
    if 'username' not in session:
        return redirect(url_for('login'))

    if request.method == 'POST':
        session.pop('username', None)
        return redirect(url_for('login'))

    return render_template('logout.html')

if __name__ == '__main__':
    app.run(host='127.0.0.1', port='5000', debug=True)


# from flask import Flask, request, render_template, send_file, redirect, url_for, session
# import os
# from docx import Document
# import mysql.connector

# app = Flask(__name__)
# app.secret_key = 'ilmy24'

# # Konfigurasi koneksi database MySQL
# db = mysql.connector.connect(
#     host="localhost",
#     user="root",
#     password="",
#     database="kriptografi"
# )

# cursor = db.cursor()

# @app.route('/', methods=['GET', 'POST'])
# def home():
#     if 'username' in session:
#         if request.method == 'POST':
#             # Handle file encryption
#             input_file = request.files['file']
#             key = request.form['key']

#             # Save the input file
#             input_file_path = os.path.join(app.root_path, 'input.txt') if input_file.filename.endswith('.txt') else os.path.join(app.root_path, 'input.docx')
#             input_file.save(input_file_path)

#             # Encrypt the file
#             encrypted_file = os.path.join(app.root_path, 'encrypted.txt') if input_file.filename.endswith('.txt') else os.path.join(app.root_path, 'encrypted.docx')
#             encrypt_file(input_file_path, encrypted_file, key)

#             # Read the input and encrypted files as binary data
#             with open(input_file_path, 'rb') as f:
#                 input_file_data = f.read()
#             with open(encrypted_file, 'rb') as f:
#                 encrypted_file_data = f.read()

#             # Store the file data in the database
#             username = session['username']
#             update_query = "UPDATE users SET input_file = %s, encrypted_file = %s WHERE username = %s"
#             cursor.execute(update_query, (input_file_data, encrypted_file_data, username))
#             db.commit()

#             # Delete the input and encrypted files
#             # # os.remove(input_file_path)
#             # # os.remove(encrypted_file)

#             return render_template('index.html', message='File encrypted successfully.', encrypted_file=encrypted_file)

#         return render_template('index.html')

#     return redirect(url_for('login'))

# @app.route('/decrypt', methods=['POST'])
# def decrypt():
#     # Handle file decryption
#     encrypted_file = request.files['file']
#     key = request.form['key']

#     # Save the encrypted file
#     encrypted_file_path = 'encrypted.txt' if encrypted_file.filename.endswith('.txt') else 'encrypted.docx'
#     encrypted_file.save(encrypted_file_path)

#     # Decrypt the file
#     decrypted_file = 'decrypted.txt' if encrypted_file.filename.endswith('.txt') else 'decrypted.docx'
#     decrypt_file(encrypted_file_path, decrypted_file, key)

#     # Delete the encrypted file
#     os.remove(encrypted_file_path)

#     return render_template('index.html', message='File decrypted successfully.', decrypted_file=decrypted_file)

# @app.route('/download/<path:filename>', methods=['GET'])
# def download(filename):
#     return send_file(filename, as_attachment=True)

# def encrypt_file(input_file, output_file, key):
#     if input_file.endswith('.txt'):
#         with open(input_file, 'r') as file:
#             text = file.read()

#         encrypted_text = encrypt_text(text, key)

#         with open(output_file, 'w') as file:
#             file.write(encrypted_text)
#     elif input_file.endswith('.docx'):
#         doc = Document(input_file)

#         for paragraph in doc.paragraphs:
#             encrypted_text = encrypt_text(paragraph.text, key)
#             paragraph.text = encrypted_text

#         doc.save(output_file)

# def decrypt_file(input_file, output_file, key):
#     if input_file.endswith('.txt'):
#         with open(input_file, 'r') as file:
#             encrypted_text = file.read()

#         decrypted_text = decrypt_text(encrypted_text, key)

#         with open(output_file, 'w') as file:
#             file.write(decrypted_text)
#     elif input_file.endswith('.docx'):
#         doc = Document(input_file)

#         for paragraph in doc.paragraphs:
#             encrypted_text = paragraph.text  # Retrieve the encrypted text
#             decrypted_text = decrypt_text(encrypted_text, key)  # Decrypt the text
#             reversed_text = reverse_text(decrypted_text)  # Reverse the decrypted text
#             paragraph.text = reversed_text  # Replace the paragraph text with the reversed text

#         doc.save(output_file)

# def encrypt_text(text, key):
#     encrypted_text = ''
#     key_index = 0
#     for char in text:
#         if char.isalpha():
#             if char.isupper():
#                 encrypted_char = chr((ord(char) - ord('A') + ord(key[key_index].upper()) - ord('A')) % 26 + ord('A'))
#             else:
#                 encrypted_char = chr((ord(char) - ord('a') + ord(key[key_index].lower()) - ord('a')) % 26 + ord('a'))

#             encrypted_text += encrypted_char

#             # Move to the next character in the key
#             key_index = (key_index + 1) % len(key)
#         else:
#             encrypted_text += char

#     return encrypted_text[::-1]  # Reverse the encrypted text


# def decrypt_text(text, key):
#     decrypted_text = ''
#     key_index = 0
#     for char in text[::-1]:  # Reverse the input text
#         if char.isalpha():
#             if char.isupper():
#                 decrypted_char = chr((ord(char) - ord('A') - (ord(key[key_index].upper()) - ord('A')) + 26) % 26 + ord('A'))
#             else:
#                 decrypted_char = chr((ord(char) - ord('a') - (ord(key[key_index].lower()) - ord('a')) + 26) % 26 + ord('a'))

#             decrypted_text += decrypted_char

#             # Move to the next character in the key
#             key_index = (key_index + 1) % len(key)
#         else:
#             decrypted_text += char

#     return decrypted_text[::-1]  # Reverse the decrypted text


# def reverse_text(text):
#     return text[::-1]

# @app.route('/login', methods=['GET', 'POST'])
# def login():
#     if 'username' in session:
#         return redirect(url_for('home'))

#     if request.method == 'POST':
#         username = request.form['username']
#         password = request.form['password']

#         # Query untuk memeriksa keberadaan username dan password di database
#         query = "SELECT * FROM users WHERE username = %s AND password = %s"
#         cursor.execute(query, (username, password))
#         user = cursor.fetchone()

#         if user:
#             session['username'] = user[1]  # Menggunakan session dengan username dari database
#             return redirect(url_for('home'))
#         else:
#             return render_template('login.html', error='Invalid username or password')

#     return render_template('login.html')

# @app.route('/logout', methods=['GET', 'POST'])
# def logout():
#     if 'username' not in session:
#         return redirect(url_for('login'))

#     if request.method == 'POST':
#         session.pop('username', None)
#         return redirect(url_for('login'))

#     return render_template('logout.html')

# if __name__ == '__main__':
#     app.run(host='127.0.0.1', port='4000', debug=True)



# # from flask import Flask, request, render_template, send_file, redirect, url_for, session
# # import os
# # from docx import Document
# # import mysql.connector

# # app = Flask(__name__)
# # app.secret_key = 'ilmy24'

# # # Konfigurasi koneksi database MySQL
# # db = mysql.connector.connect(
# #     host="localhost",
# #     user="root",
# #     password="",
# #     database="kriptografi"
# # )

# # cursor = db.cursor()

# # @app.route('/', methods=['GET', 'POST'])
# # def home():
# #     if 'username' in session:
# #         if request.method == 'POST':
# #             # Handle file encryption
# #             input_file = request.files['file']
# #             key = request.form['key']

# #             # Save the input file
# #             input_file_path = os.path.join(app.root_path, 'input.txt') if input_file.filename.endswith('.txt') else os.path.join(app.root_path, 'input.docx')
# #             input_file.save(input_file_path)

# #             # Encrypt the file
# #             encrypted_file = os.path.join(app.root_path, 'encrypted.txt') if input_file.filename.endswith('.txt') else os.path.join(app.root_path, 'encrypted.docx')
# #             encrypt_file(input_file_path, encrypted_file, key)

# #             # Read the input and encrypted files as binary data
# #             with open(input_file_path, 'rb') as f:
# #                 input_file_data = f.read()
# #             with open(encrypted_file, 'rb') as f:
# #                 encrypted_file_data = f.read()

# #             # Store the file data in the database
# #             username = session['username']
# #             update_query = "UPDATE users SET input_file = %s, encrypted_file = %s WHERE username = %s"
# #             cursor.execute(update_query, (input_file_data, encrypted_file_data, username))
# #             db.commit()

# #             # Delete the input and encrypted files
# #             # # os.remove(input_file_path)
# #             # # os.remove(encrypted_file)

# #             return render_template('index.html', message='File encrypted successfully.', encrypted_file=encrypted_file)

# #         return render_template('index.html')

# #     return redirect(url_for('login'))

# # @app.route('/decrypt', methods=['POST'])
# # def decrypt():
# #     # Handle file decryption
# #     encrypted_file = request.files['file']
# #     key = request.form['key']

# #     # Save the encrypted file
# #     encrypted_file_path = 'encrypted.txt' if encrypted_file.filename.endswith('.txt') else 'encrypted.docx'
# #     encrypted_file.save(encrypted_file_path)

# #     # Decrypt the file
# #     decrypted_file = 'decrypted.txt' if encrypted_file.filename.endswith('.txt') else 'decrypted.docx'
# #     decrypt_file(encrypted_file_path, decrypted_file, key)

# #     # Delete the encrypted file
# #     os.remove(encrypted_file_path)

# #     return render_template('index.html', message='File decrypted successfully.', decrypted_file=decrypted_file)

# # @app.route('/download/<path:filename>', methods=['GET'])
# # def download(filename):
# #     return send_file(filename, as_attachment=True)

# # def encrypt_file(input_file, output_file, key):
# #     if input_file.endswith('.txt'):
# #         with open(input_file, 'r') as file:
# #             text = file.read()

# #         encrypted_text = encrypt_text(text, key)

# #         with open(output_file, 'w') as file:
# #             file.write(encrypted_text)
# #     elif input_file.endswith('.docx'):
# #         doc = Document(input_file)

# #         for paragraph in doc.paragraphs:
# #             encrypted_text = encrypt_text(paragraph.text, key)
# #             paragraph.text = encrypted_text

# #         doc.save(output_file)

# # def decrypt_file(input_file, output_file, key):
# #     if input_file.endswith('.txt'):
# #         with open(input_file, 'r') as file:
# #             encrypted_text = file.read()

# #         decrypted_text = decrypt_text(encrypted_text, key)

# #         with open(output_file, 'w') as file:
# #             file.write(decrypted_text)
# #     elif input_file.endswith('.docx'):
# #         doc = Document(input_file)

# #         for paragraph in doc.paragraphs:
# #             decrypted_text = decrypt_text(paragraph.text, key)
# #             paragraph.text = decrypted_text

# #         doc.save(output_file)

# # def encrypt_text(text, key):
# #     encrypted_text = ''
# #     key_index = 0
# #     for char in text:
# #         if char.isalpha():
# #             if char.isupper():
# #                 encrypted_char = chr((ord(char) - ord('A') + ord(key[key_index].upper()) - ord('A')) % 26 + ord('A'))
# #             else:
# #                 encrypted_char = chr((ord(char) - ord('a') + ord(key[key_index].lower()) - ord('a')) % 26 + ord('a'))

# #             encrypted_text += encrypted_char

# #             # Move to the next character in the key
# #             key_index = (key_index + 1) % len(key)
# #         else:
# #             encrypted_text += char

# #     return encrypted_text

# # def decrypt_text(text, key):
# #     decrypted_text = ''
# #     key_index = 0
# #     for char in text:
# #         if char.isalpha():
# #             if char.isupper():
# #                 decrypted_char = chr((ord(char) - ord('A') - (ord(key[key_index].upper()) - ord('A')) + 26) % 26 + ord('A'))
# #             else:
# #                 decrypted_char = chr((ord(char) - ord('a') - (ord(key[key_index].lower()) - ord('a')) + 26) % 26 + ord('a'))

# #             decrypted_text += decrypted_char

# #             # Move to the next character in the key
# #             key_index = (key_index + 1) % len(key)
# #         else:
# #             decrypted_text += char

# #     return decrypted_text

# # @app.route('/login', methods=['GET', 'POST'])
# # def login():
# #     if 'username' in session:
# #         return redirect(url_for('home'))

# #     if request.method == 'POST':
# #         username = request.form['username']
# #         password = request.form['password']

# #         # Query untuk memeriksa keberadaan username dan password di database
# #         query = "SELECT * FROM users WHERE username = %s AND password = %s"
# #         cursor.execute(query, (username, password))
# #         user = cursor.fetchone()

# #         if user:
# #             session['username'] = user[1]  # Menggunakan session dengan username dari database
# #             return redirect(url_for('home'))
# #         else:
# #             return render_template('login.html', error='Invalid username or password')

# #     return render_template('login.html')

# # @app.route('/logout', methods=['GET', 'POST'])
# # def logout():
# #     if 'username' not in session:
# #         return redirect(url_for('login'))

# #     if request.method == 'POST':
# #         session.pop('username', None)
# #         return redirect(url_for('login'))

# #     return render_template('logout.html')

# # if __name__ == '__main__':
# #     app.run(host='127.0.0.1', port='5000', debug=True)
